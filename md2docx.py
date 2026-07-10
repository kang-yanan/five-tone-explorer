"""
Convert paper_final.md to Nature-style DOCX.
Features: three-line tables, Arial font, preserved bold, column alignment.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SRC = "F:/Claude project/five_tone_experiment/paper_ubicomp.md"
DST = "F:/Claude project/five_tone_experiment/paper_ubicomp_v2.docx"
IMG_DIR = "F:/Claude project/five_tone_experiment"

def set_cell_font(cell, text, bold=False, size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Set cell text with Arial font and optional bold."""
    # Clear existing paragraphs
    for p in cell.paragraphs:
        for r in p.runs:
            r.clear()
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)

def make_three_line_table(table):
    """Apply Nature-style three-line borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # Remove existing borders element if present
    for child in list(tblPr):
        if child.tag == qn('w:tblBorders'):
            tblPr.remove(child)

    # Three-line border spec: thick top, thin under header, thick bottom
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Thin border below header row (row 0)
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        # Remove existing cell borders
        for child in list(tcPr):
            if child.tag == qn('w:tcBorders'):
                tcPr.remove(child)
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

def parse_table_cell(cell_text):
    """Parse cell text: return (text, is_bold) for each segment.
    Supports **bold** markers."""
    parts = re.split(r'(\*\*(.+?)\*\*)', cell_text)
    # parts alternates: plain, full_marker, bold_content, plain...
    result = []
    for idx, part in enumerate(parts):
        if idx % 3 == 0:  # plain text
            if part.strip():
                result.append((part, False))
        elif idx % 3 == 2:  # bold content (skip the full marker at idx%3==1)
            result.append((part, True))
    # If no bold markers found, return whole text as plain
    if not result:
        return [(cell_text, False)]
    return result

def render_nature_table(doc, tlines):
    """Render a Markdown table as a Nature-style three-line table."""
    if len(tlines) < 2:
        return

    # Parse header alignment directives
    hdr_raw = [c.strip() for c in tlines[0].split("|") if c.strip()]
    sep_raw = [c.strip() for c in tlines[1].split("|") if c.strip()]

    # Determine column alignments from separator row
    alignments = []
    for s in sep_raw:
        s = s.replace(":", "").strip("-")
        alignments.append(WD_ALIGN_PARAGRAPH.CENTER)  # all centered by default

    # More precise: check :---, :---:, ---:
    seps = [c.strip() for c in tlines[1].split("|") if c.strip()]
    for idx, s in enumerate(seps):
        left = s.startswith(":")
        right = s.endswith(":")
        if left and right:
            alignments[idx] = WD_ALIGN_PARAGRAPH.CENTER
        elif right:
            alignments[idx] = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            alignments[idx] = WD_ALIGN_PARAGRAPH.LEFT

    # Parse data rows (skip separator line at index 1)
    data_rows = []
    for rline in tlines[2:]:
        cells = [c.strip() for c in rline.split("|") if c.strip()]
        data_rows.append(cells)

    ncols = len(hdr_raw)
    nrows = 1 + len(data_rows)
    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row (Arial 9pt bold, centered)
    for j, h in enumerate(hdr_raw):
        set_cell_font(tbl.rows[0].cells[j], h, bold=True, size=Pt(9),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, row in enumerate(data_rows):
        for ci, cell_text in enumerate(row):
            if ci >= ncols:
                continue
            # Parse bold markers
            segments = parse_table_cell(cell_text)
            cell = tbl.rows[ri + 1].cells[ci]
            # Clear default paragraph
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = alignments[ci] if ci < len(alignments) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            for text, is_bold in segments:
                run = p.add_run(text)
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.font.bold = is_bold
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Apply three-line borders
    make_three_line_table(tbl)

    # Add a small space after the table
    doc.add_paragraph("")


def main():
    with open(SRC, "r", encoding="utf-8") as f:
        md = f.read()

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]; i += 1
        if not line.strip():
            continue

        # Headings
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:], level=0)
            for run in p.runs:
                run.font.name = "Arial"
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=1)
            for run in p.runs:
                run.font.name = "Arial"
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=2)
            for run in p.runs:
                run.font.name = "Arial"

        # Centered bold (author / affiliation / draft date)
        elif line.strip().startswith("**") and line.strip().endswith("**"):
            p = doc.add_paragraph(line.strip("*"))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Arial"

        # Italic centered
        elif line.startswith("*") and not line.startswith("**") and len(line.strip()) < 80:
            p = doc.add_paragraph(line.strip("*"))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Arial"
                run.font.italic = True

        # Tables — Nature-style three-line
        elif line.startswith("|"):
            tlines = [line]
            while i < len(lines) and lines[i].startswith("|"):
                tlines.append(lines[i]); i += 1
            render_nature_table(doc, tlines)

        # Images
        elif line.startswith("!["):
            alt = line[line.index("[")+1:line.index("]")]
            src = line[line.index("(")+1:line.index(")")]
            try:
                doc.add_picture(IMG_DIR + "/" + src, width=Inches(5.2))
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.name = "Arial"
            except Exception:
                p = doc.add_paragraph("[Figure: " + src + "]")
                for run in p.runs:
                    run.font.name = "Arial"

        # Bullet / numbered lists
        elif line.startswith("- ") or (len(line) > 2 and line[0].isdigit() and ". " in line[:4]):
            items = [re.sub(r"^\d+\.\s*", "", line).lstrip("- ")]
            while i < len(lines) and (
                lines[i].startswith("- ") or
                (len(lines[i]) > 2 and lines[i][0].isdigit() and ". " in lines[i][:4])
            ):
                item = re.sub(r"^\d+\.\s*", "", lines[i]).lstrip("- ")
                items.append(item)
                i += 1
            for item in items:
                p = doc.add_paragraph(item, style="List Bullet")
                for run in p.runs:
                    run.font.name = "Arial"

        # Separators and HTML tags
        elif line == "---" or line.startswith("<"):
            continue

        # Regular paragraph
        else:
            clean = line
            # Handle bold: preserve as formatting
            bold_segments = re.split(r'(\*\*(.+?)\*\*)', clean)
            p = doc.add_paragraph()
            for seg_idx in range(0, len(bold_segments)):
                seg = bold_segments[seg_idx]
                if not seg:
                    continue
                if seg_idx % 3 == 1:  # **full marker**, skip (content at seg_idx+1)
                    continue
                if seg_idx % 3 == 2:  # bold content
                    run = p.add_run(seg)
                    run.font.bold = True
                else:  # plain text
                    run = p.add_run(seg)
                run.font.name = "Arial"
                run.font.size = Pt(11)

            # Handle italic
            for run in p.runs:
                if run.text.startswith("*") and run.text.endswith("*") and len(run.text) > 2:
                    run.text = run.text[1:-1]
                    run.font.italic = True

    doc.save(DST)
    print("Nature-style Word document saved:", DST)

if __name__ == "__main__":
    main()
